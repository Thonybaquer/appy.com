import pandas as pd
from docx import Document

def exportar_a_excel(df, nombre_archivo):
    """Exporta un DataFrame a un archivo Excel."""
    df.to_excel(nombre_archivo + '.xlsx', index=False)

def exportar_a_pdf(df, nombre_archivo):
    """Exporta un DataFrame a un archivo PDF."""
    # Lógica para exportar a PDF
    pass

def exportar_a_csv(df, nombre_archivo):
    """Exporta un DataFrame a un archivo CSV."""
    df.to_csv(nombre_archivo + '.csv', index=False)

def exportar_a_json(df, nombre_archivo):
    """Exporta un DataFrame a un archivo JSON."""
    df.to_json(nombre_archivo + '.json', orient='records')

def exportar_a_txt(df, nombre_archivo):
    """Exporta un DataFrame a un archivo TXT."""
    df.to_csv(nombre_archivo + '.txt', index=False, sep='\t')

def exportar_a_html(df, nombre_archivo):
    """Exporta un DataFrame a un archivo HTML."""
    df.to_html(nombre_archivo + '.html', index=False)

def exportar_a_word(df, nombre_archivo):
    """Exporta un DataFrame a un archivo Word."""
    doc = Document()
    doc.add_heading('Datos Exportados', level=1)

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.save(nombre_archivo + '.docx') 